"""
Enhanced PDF to Word converter using PyMuPDF.
Better quality than pdf2docx for complex PDFs.
"""
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re


def extract_pdf_content_pymupdf(pdf_path):
    """
    Extract structured content from PDF using PyMuPDF.
    Returns dict with text blocks, tables, and metadata.
    """
    doc = fitz.open(pdf_path)
    content = {
        "text_blocks": [],
        "tables": [],
        "metadata": doc.metadata,
        "num_pages": len(doc)
    }
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        # Extract text blocks with font info
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if "lines" in block:  # Text block
                for line in block["lines"]:
                    text = ""
                    font_size = 12
                    is_bold = False
                    
                    for span in line["spans"]:
                        text += span["text"]
                        font_size = span["size"]
                        is_bold = "bold" in span["font"].lower()
                    
                    if text.strip():
                        content["text_blocks"].append({
                            "text": text.strip(),
                            "font_size": font_size,
                            "is_bold": is_bold,
                            "page": page_num + 1
                        })
    
    doc.close()
    return content


def create_word_document(content, output_path):
    """
    Create Word document from extracted PDF content.
    """
    doc = Document()
    
    # Process text blocks
    for block in content["text_blocks"]:
        text = block["text"]
        font_size = block["font_size"]
        is_bold = block["is_bold"]
        
        # Detect headings (larger font or bold)
        if font_size > 14 or (is_bold and len(text) < 100):
            # Treat as heading
            paragraph = doc.add_heading(text, level=1 if font_size > 16 else 2)
        else:
            # Regular paragraph
            paragraph = doc.add_paragraph(text)
            
            # Apply formatting
            run = paragraph.runs[0] if paragraph.runs else None
            if run:
                run.font.size = Pt(min(font_size, 12))  # Cap at 12pt
                if is_bold:
                    run.font.bold = True
    
    doc.save(output_path)
    return doc


def convert_pdf_to_word_pymupdf(pdf_path, output_path):
    """
    Main conversion function using PyMuPDF.
    """
    try:
        content = extract_pdf_content_pymupdf(pdf_path)
        doc = create_word_document(content, output_path)
        return True, "pymupdf"
    except Exception as e:
        print(f"[ERROR] PyMuPDF conversion failed: {e}")
        return False, str(e)
