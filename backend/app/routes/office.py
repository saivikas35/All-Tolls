"""
Office document → PDF conversion using LibreOffice.
Supports: Word (.docx/.doc), PowerPoint (.pptx/.ppt), Excel (.xlsx/.xls)
Falls back to a helpful error if LibreOffice is not installed.
"""
from fastapi import APIRouter, UploadFile, File, HTTPException, Form
import subprocess
import os
import uuid
import shutil
import platform
import copy
from typing import Optional, List
from app.utils import get_file_or_url, save_upload, generate_download_url, validate_output_file

router = APIRouter()
UPLOAD_DIR = "uploads"


def get_soffice_path():
    """Find LibreOffice executable path."""
    env_path = os.getenv("SOFFICE_PATH")
    if env_path and os.path.exists(env_path):
        return env_path

    if platform.system() == "Windows":
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            r"C:\Program Files\LibreOffice 7\program\soffice.exe",
            r"C:\Program Files\LibreOffice 6\program\soffice.exe",
        ]
        for path in possible_paths:
            if os.path.exists(path):
                return path
        return None  # Not found

    # Linux/Mac
    return shutil.which("soffice") or shutil.which("libreoffice")


def _convert_to_pdf_via_libreoffice(input_path: str) -> str:
    """Run LibreOffice headless to convert to PDF. Returns output PDF path."""
    soffice = get_soffice_path()
    if not soffice:
        raise HTTPException(
            status_code=503,
            detail=(
                "LibreOffice is not installed on this server. "
                "Please install LibreOffice to use Office-to-PDF conversions. "
                "Download from: https://www.libreoffice.org/download/download-libreoffice/"
            ),
        )

    out_dir = os.path.abspath(UPLOAD_DIR)
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", out_dir,
            os.path.abspath(input_path),
        ],
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise HTTPException(
            status_code=500,
            detail=f"LibreOffice conversion failed: {result.stderr or result.stdout}"
        )

    # LibreOffice outputs <input_basename>.pdf in out_dir
    base = os.path.splitext(os.path.basename(input_path))[0]
    out_pdf = os.path.join(out_dir, f"{base}.pdf")
    if not os.path.exists(out_pdf):
        raise HTTPException(status_code=500, detail="Conversion produced no output file.")

    # Rename with unique identifier
    unique_name = f"{uuid.uuid4().hex[:8]}_{base}_AllTools.pdf"
    unique_path = os.path.join(out_dir, unique_name)
    os.rename(out_pdf, unique_path)
    return unique_name


def _prepare_excel_for_pdf(input_path: str):
    """Modify Excel to fit to 1 page wide for better PDF output."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(input_path)
        for sheet in wb.worksheets:
            sheet.page_setup.fitToWidth = 1
            sheet.page_setup.fitToHeight = 0  # 0 means automatic
        wb.save(input_path)
    except Exception as e:
        print(f"Warning: Could not format Excel file with openpyxl: {e}")


# ─── WORD → PDF ───────────────────────────────────────────────────────────────

@router.post("/word-to-pdf")
def word_to_pdf(
    file: Optional[UploadFile] = File(default=None),
    url: Optional[str] = Form(default=None)
):
    """Convert Word document (.docx/.doc) to PDF."""
    allowed = (".docx", ".doc", ".odt", ".rtf")
    if file and file.filename:
        if not any(file.filename.lower().endswith(ext) for ext in allowed):
            raise HTTPException(status_code=400, detail="Supported formats: .docx, .doc, .odt, .rtf")
        ext = os.path.splitext(file.filename)[1].lower()
    else:
        # If url, assume docx as fallback extension
        ext = ".docx"

    in_path = get_file_or_url(file, url, suffix=ext)
    out_name = _convert_to_pdf_via_libreoffice(in_path)
    validate_output_file(os.path.join(UPLOAD_DIR, out_name), ".pdf")
    return {"success": True, "downloadUrl": generate_download_url(out_name)}


# ─── POWERPOINT → PDF ─────────────────────────────────────────────────────────

@router.post("/ppt-to-pdf")
def ppt_to_pdf(
    file: Optional[UploadFile] = File(default=None),
    url: Optional[str] = Form(default=None)
):
    """Convert PowerPoint presentation (.pptx/.ppt) to PDF."""
    allowed = (".pptx", ".ppt", ".odp")
    if file and file.filename:
        if not any(file.filename.lower().endswith(ext) for ext in allowed):
            raise HTTPException(status_code=400, detail="Supported formats: .pptx, .ppt, .odp")
        ext = os.path.splitext(file.filename)[1].lower()
    else:
        ext = ".pptx"

    in_path = get_file_or_url(file, url, suffix=ext)
    out_name = _convert_to_pdf_via_libreoffice(in_path)
    validate_output_file(os.path.join(UPLOAD_DIR, out_name), ".pdf")
    return {"success": True, "downloadUrl": generate_download_url(out_name)}


# ─── EXCEL → PDF ──────────────────────────────────────────────────────────────

@router.post("/excel-to-pdf")
def excel_to_pdf(
    file: Optional[UploadFile] = File(default=None),
    url: Optional[str] = Form(default=None)
):
    """Convert Excel spreadsheet (.xlsx/.xls) to PDF."""
    allowed = (".xlsx", ".xls", ".ods", ".csv")
    if file and file.filename:
        if not any(file.filename.lower().endswith(ext) for ext in allowed):
            raise HTTPException(status_code=400, detail="Supported formats: .xlsx, .xls, .ods, .csv")
        ext = os.path.splitext(file.filename)[1].lower()
    else:
        ext = ".xlsx"

    in_path = get_file_or_url(file, url, suffix=ext)
    if ext == ".xlsx":
        _prepare_excel_for_pdf(in_path)
    out_name = _convert_to_pdf_via_libreoffice(in_path)
    validate_output_file(os.path.join(UPLOAD_DIR, out_name), ".pdf")
    return {"success": True, "downloadUrl": generate_download_url(out_name)}


# ─── WORD → PDF PREVIEW (for page-selector thumbnail generation) ──────────────

@router.post("/word-to-pdf-preview")
def word_to_pdf_preview(
    file: Optional[UploadFile] = File(default=None),
    url: Optional[str] = Form(default=None)
):
    """Convert Word doc to PDF and return the PDF token so the frontend can
    call /api/convert/pdf-split/preview to render page thumbnails."""
    allowed = (".docx", ".doc", ".odt", ".rtf")
    if file and file.filename:
        if not any(file.filename.lower().endswith(ext) for ext in allowed):
            raise HTTPException(status_code=400, detail="Supported formats: .docx, .doc, .odt, .rtf")
        ext = os.path.splitext(file.filename)[1].lower()
    else:
        ext = ".docx"

    in_path = get_file_or_url(file, url, suffix=ext)
    # Save original docx path for later removal
    docx_path = in_path
    pdf_name = _convert_to_pdf_via_libreoffice(in_path)
    return {
        "success": True,
        "pdfToken": pdf_name,
        "docxToken": os.path.basename(docx_path),
    }


# ─── WORD REMOVE PAGES ────────────────────────────────────────────────────────

def _split_docx_into_pages(doc):
    """Split a python-docx Document into page-groups by explicit page breaks.
    Returns a list of lists of paragraphs (each inner list = one 'page')."""
    from docx.enum.text import WD_BREAK
    pages = []
    current_page = []

    for para in doc.paragraphs:
        current_page.append(para)
        # Check runs for page-break
        has_page_break = False
        for run in para.runs:
            if run.contains_page_break:
                has_page_break = True
                break
            # Also check XML directly
            for br in run._r.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'):
                btype = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', '')
                if btype == 'page':
                    has_page_break = True
                    break

        if has_page_break:
            pages.append(current_page)
            current_page = []

    if current_page:
        pages.append(current_page)

    return pages


@router.post("/word-remove-pages")
def word_remove_pages(
    file: Optional[UploadFile] = File(default=None),
    url: Optional[str] = Form(default=None),
    docx_token: Optional[str] = Form(default=None),
    pages: str = Form(default=""),
    output_format: str = Form(default="docx"),
):
    """Remove selected pages from a Word document.
    
    - pages: comma-separated 1-indexed page numbers to REMOVE
    - output_format: 'docx' or 'pdf'
    - docx_token: optional cached .docx filename from word-to-pdf-preview step
    """
    import re
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from lxml import etree

    pages = (pages or "").strip()
    output_format = (output_format or "docx").strip().lower()

    if not pages:
        raise HTTPException(status_code=400, detail="Please specify pages to remove.")

    # Resolve input path
    if docx_token:
        in_path = os.path.join(UPLOAD_DIR, os.path.basename(docx_token))
        if not os.path.exists(in_path):
            raise HTTPException(status_code=400, detail="Session expired. Please re-upload your file.")
    else:
        ext = ".docx"
        if file and file.filename:
            if not any(file.filename.lower().endswith(e) for e in (".docx", ".doc", ".odt", ".rtf")):
                raise HTTPException(status_code=400, detail="Supported formats: .docx, .doc, .odt, .rtf")
            ext = os.path.splitext(file.filename)[1].lower()
        in_path = get_file_or_url(file, url, suffix=ext)

    # Parse page numbers to remove (1-indexed)
    remove_nums = set()
    for part in pages.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            for n in range(int(a), int(b) + 1):
                remove_nums.add(n)
        elif part.isdigit():
            remove_nums.add(int(part))

    # Load the document
    doc = Document(in_path)
    page_groups = _split_docx_into_pages(doc)
    total_pages = len(page_groups)

    if not remove_nums - {0}:
        raise HTTPException(status_code=400, detail="No valid pages specified.")
    if len(remove_nums) >= total_pages:
        raise HTTPException(status_code=400, detail="Cannot remove all pages from the document.")

    # Build list of paragraphs to KEEP (pages not in remove set)
    kept_paras = []
    for page_idx, para_list in enumerate(page_groups):
        page_num = page_idx + 1  # 1-indexed
        if page_num not in remove_nums:
            kept_paras.extend(para_list)

    # Build a new document with only the kept paragraphs
    new_doc = Document()
    # Remove default empty paragraph
    for p in new_doc.paragraphs:
        p._element.getparent().remove(p._element)

    for para in kept_paras:
        # Deep-copy the paragraph XML into the new document body
        new_para = copy.deepcopy(para._element)
        new_doc.element.body.append(new_para)

    # Copy styles from original
    try:
        new_doc.styles = doc.styles
    except Exception:
        pass

    out_docx_name = f"removed_{uuid.uuid4().hex[:8]}_AllTools.docx"
    out_docx_path = os.path.join(UPLOAD_DIR, out_docx_name)
    new_doc.save(out_docx_path)

    if output_format == "pdf":
        out_pdf_name = _convert_to_pdf_via_libreoffice(out_docx_path)
        validate_output_file(os.path.join(UPLOAD_DIR, out_pdf_name), ".pdf")
        return {"success": True, "downloadUrl": generate_download_url(out_pdf_name), "format": "pdf"}

    validate_output_file(out_docx_path, ".docx")
    return {"success": True, "downloadUrl": generate_download_url(out_docx_name), "format": "docx"}
